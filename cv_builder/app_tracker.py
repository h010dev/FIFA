import os
import docx
import datetime
import pandas as pd
import sqlite3
from my_projects.cv_builder.expense_range import df_index_table
from my_projects.cv_builder.us_states import states
from my_projects.cv_builder.secret_vars import *


pd.set_option('display.max_columns', 20)
pd.set_option('display.width', 1000)


class JobTracker:

    key_words = dict()

    def __init__(self, df=None, cost_of_living_index=df_index_table, us_states_table=states, document=None,
                 database=None, conn=None, c=None):

        """
        Initiate base parameters/variables.
        :param df: data frame to read csv file to
        :param cost_of_living_index: data frame with various cost of living indices for cities around the globe
        :param us_states_table: data frame with US states and their associated abbreviations
        :param document: cover letter template to be later edited
        :param database: sqlite3 database to read/write job applications to
        :param conn: connection to sqlite3
        :param c: cursor for sqlite3 connection
        """

        self.df = df
        self.cost_of_living_index = cost_of_living_index
        self.us_states_table = us_states_table
        self.document = document
        self.database = database
        self.conn = conn
        self.c = c

    def reload_files(self):

        """
        :return: Load initial files and subsequently reload any changes made to them.
        """

        self.df = pd.read_csv(r'my_projects\cv_builder\data\app_tracker.csv')
        self.document = docx.Document(r'my_projects\cv_builder\data\cv_template.docx')
        self.database = r'my_projects\cv_builder\data\app_tracker.db'
        self.conn = sqlite3.connect(self.database)
        self.df.tail()
        self.c = self.conn.cursor()
        self.total_apps_today()

    def organization_matched(self, org):

        """
        :param org: name of organization to search for
        :return: return sqlite3 query for matches made with org
        """

        data = self.c.execute(f"SELECT Organization, Position, ProvinceState, ApplicationDate "
                              f"FROM Applications WHERE Organization LIKE '{org}'")
        for row in data:
            print(row)

    def calculate_living_expenses(self):

        """
        :return: Calculate monthly and annual expenses, based on current expenses relative to cost of living index in
        specified state.
        """

        if not self.cost_of_living_index[
                self.cost_of_living_index['state'] == self.key_words['State Abbreviation']].empty:
            monthly_expenses_usd = \
                self.cost_of_living_index[self.cost_of_living_index['country'] == 'United States']\
                    .groupby(['state'])['min_monthly_budget_usd'].mean()[self.key_words['State Abbreviation']]
            annual_expenses_usd = \
                self.cost_of_living_index[self.cost_of_living_index['country'] == 'United States']\
                    .groupby(['state'])['min_annual_budget_usd'].mean()[self.key_words['State Abbreviation']]

        elif self.cost_of_living_index[
                self.cost_of_living_index['state'] == self.key_words['State Abbreviation']].empty:
            monthly_expenses_usd = \
                self.cost_of_living_index[self.cost_of_living_index['country'] == 'United States'][
                    'min_monthly_budget_usd'].mean()
            annual_expenses_usd = self.cost_of_living_index[
                self.cost_of_living_index['country'] == 'United States']['min_annual_budget_usd'].mean()

        return monthly_expenses_usd, annual_expenses_usd

    def new_application(self):

        """
        :return: Asks user to input information regarding new application. Prints out any matching organizations with
        application details, as well as a summary of living expenses for the input location.
        """

        position = str(input('Enter position name: '))
        organization = str(input('Enter organization name: '))
        location = str(input('Enter location: '))

        self.key_words.clear()
        self.key_words['Job'] = position
        self.key_words['Organization'] = organization
        self.key_words['Location'] = location

        self.key_words['State Abbreviation'] = self.us_states_table[
            self.us_states_table['State'] == location].reset_index()['Abbreviation'][0]

        self.organization_matched(organization)

        monthly, annual = self.calculate_living_expenses()

        print(f'Monthly Min: ${monthly:.2f}\nAnnual Min: ${annual:.2f}')

    def create_cover_letter(self):

        """
        :return: Saves a copy of the cover letter template after searching for and replacing key-words matching
        pre-defined variables.
        """

        for i in range(len(self.document.paragraphs)):
            self.document.paragraphs[i].text = self.document.paragraphs[i].text.replace('JOB', self.key_words['Job'])
            self.document.paragraphs[i].text = \
                self.document.paragraphs[i].text.replace('ORGANIZATION', self.key_words['Organization'])
            self.document.paragraphs[i].text = \
                self.document.paragraphs[i].text.replace('LOCATION', self.key_words['Location'])

        file_name = f'{self.key_words["Organization"].replace(" ", "-")}' \
                    f'_CV_{self.key_words["Job"].replace(" ", "-")}' \
                    f'_{self.key_words["Location"].replace(" ", "-")}' \
                    f'_USA_{datetime.date.today()}.docx'

        save_path = os.path.join(file_path, file_name).replace('\\', '/')

        self.document.save(save_path)

    def update_csv(self):

        """
        :return: Updates csv file with new job application details.
        """

        d = dict()

        d['Organization'] = self.key_words['Organization']
        d['Position'] = self.key_words['Job']
        d['Province/State'] = self.key_words['Location']
        d['Country'] = 'USA'
        d['Application Date'] = datetime.date.today()
        d['Test'] = False
        d['Contact'] = False
        d['Interview'] = False
        d['Notes'] = str(input('Enter any notes here: '))
        d['Rejected'] = False

        self.df = self.df.append(d, ignore_index=True)
        self.df.to_csv(r'my_projects\cv_builder\data\app_tracker.csv', index=False)
        d.clear()

    def total_apps_today(self):

        """
        :return: Number of applications done today.
        """

        data = self.c.execute(f"SELECT * FROM Applications "
                              f"WHERE ApplicationDate "
                              f"BETWEEN '{str(datetime.date.today())} 00:00:00' "
                              f"AND '{str(datetime.date.today())} 23:59:59'")

        count = 0

        for row in data:
            print(row)
            count += 1

        print(f'Jobs Applied to Today: {count}')

    def create_sql_table(self):

        """
        :return: Create new sqlite3 table if none exists.
        """

        self.c.execute("CREATE TABLE IF NOT EXISTS Applications(Organization TEXT, Position TEXT, ProvinceState TEXT, "
                       "Country TEXT, ApplicationDate TEXT, Test INTEGER, Contact INTEGER, Interview INTEGER, "
                       "Rejected INTEGER, Notes TEXT)")

    def update_sql(self):

        """
        :return: Updates sqlite3 database with new job application details.
        """

        organization = self.key_words['Organization']
        position = self.key_words['Job']
        province_state = self.key_words['Location']
        country = 'USA'
        application_date = str(datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        test = 0
        contact = 0
        interview = 0
        rejected = 0
        notes = str(input('Enter any notes here: '))

        self.c.execute("INSERT INTO Applications (Organization, Position, ProvinceState, Country, ApplicationDate, "
                       "Test, Contact, Interview, Rejected, Notes) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                       (organization, position, province_state, country, application_date, test, contact, interview,
                        rejected, notes))

        self.conn.commit()
        self.conn.close()
        self.reload_files()

